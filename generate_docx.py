import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top','left','bottom','right']:
        el = OxmlElement('w:%s' % edge)
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'4')
        el.set(qn('w:color'),'888888')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, rgb=(0x1F,0x4E,0x79)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(*rgb)
    return h

def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def tbl_header(table, headers, bg='1F4E79'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, bg)
        set_borders(c)

def fill_row(row, values, alt=None):
    for i, v in enumerate(values):
        c = row.cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(str(v))
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        set_borders(c)
        if alt:
            set_cell_bg(c, alt)

# =========================================================
#  COVER PAGE
# =========================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TRUONG DAI HOC CONG NGHE THONG TIN - KHOA CNPM')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.bold=True
r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TAI LIEU THIET KE LOP VA CODE SKELETON')
r.font.name='Times New Roman'; r.font.size=Pt(16); r.bold=True
r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Class Diagram & Java Code Skeleton - Tuan 3')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.bold=True

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HE THONG QUAN LY DAT DO AN TRUC TUYEN')
r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
r.font.color.rgb = RGBColor(0xC0,0x00,0x00)

doc.add_paragraph()

ct = doc.add_table(rows=7, cols=2)
ct.style = 'Table Grid'
cover_data = [
    ('Nhom:', 'TEAM 12 - He thong Dat do an Truc tuyen'),
    ('Thanh vien:', 'Thanh vien 1 (MSSV1)'),
    ('', 'Thanh vien 2 (MSSV2)'),
    ('', 'Thanh vien 3 (MSSV3)'),
    ('Tuan:', 'Tuan 3 - Thiet ke Lop & Code Skeleton'),
    ('GitHub:', 'https://github.com/<USERNAME>/PTTKPM25_26_NXX_Nhom12'),
    ('Phien ban:', '1.0 - Thang 4/2026'),
]
for ri,(k,v) in enumerate(cover_data):
    row = ct.rows[ri]
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in row.cells:
        for rn in cell.paragraphs[0].runs:
            rn.font.name='Times New Roman'; rn.font.size=Pt(11)
    if k:
        for rn in row.cells[0].paragraphs[0].runs: rn.bold=True
    set_borders(row.cells[0]); set_borders(row.cells[1])
doc.add_page_break()

# =========================================================
#  1. NOUN EXTRACTION
# =========================================================
heading(doc, '1. TRICH XUAT LOP TU KICH BAN (NOUN EXTRACTION)', 1)
para(doc, 'Phuong phap: Doc lai tat ca kich ban use case, gach chan danh tu, loai bo danh tu trung lap va danh tu khong phai doi tuong nghiep vu.')

nt = doc.add_table(rows=1, cols=3)
nt.style = 'Table Grid'
tbl_header(nt, ['Danh tu tim duoc', 'Phan loai', 'Ket luan'])

# Lớp thực tế theo code: User, CustomerProfile, RestaurantProfile, DriverProfile,
# Category, MenuItem, FoodOrder, OrderItem, Payment, Review, Voucher
noun_data = [
    ('Nguoi dung',           'Lop doi tuong', '-> User (Entity)  --  lop co so, co truong role'),
    ('Khach hang',           'Lop doi tuong', '-> CustomerProfile (Entity)  --  @OneToOne voi User'),
    ('Nha hang',             'Lop doi tuong', '-> RestaurantProfile (Entity)  --  @OneToOne voi User'),
    ('Tai xe / Shipper',     'Lop doi tuong', '-> DriverProfile (Entity)  --  @OneToOne voi User'),
    ('Don hang',             'Lop doi tuong', '-> FoodOrder (Entity)  --  trung tam he thong'),
    ('Danh muc mon an',      'Lop doi tuong', '-> Category (Entity)'),
    ('Mon an',               'Lop doi tuong', '-> MenuItem (Entity)'),
    ('Dong item don hang',   'Lop doi tuong', '-> OrderItem (Entity)  --  snapshot gia'),
    ('Giao dich thanh toan', 'Lop doi tuong', '-> Payment (Entity)'),
    ('Danh gia',             'Lop doi tuong', '-> Review (Entity)'),
    ('Ma khuyen mai',        'Lop doi tuong', '-> Voucher (Entity)'),
    ('Mat khau / Token',     'Thuoc tinh',    '-> Thuoc tinh trong User, khong tao lop rieng'),
    ('Tong tien / Gia',      'Thuoc tinh',    '-> Thuoc tinh Double trong FoodOrder / MenuItem'),
    ('He thong thanh toan',  'Actor phu',     '-> PaymentService interface, khong phai Entity'),
    ('Thong bao he thong',   'Actor phu',     '-> NotificationService interface, khong phai Entity'),
]
for i, rd in enumerate(noun_data):
    fill_row(nt.add_row(), rd, alt='F2F2F2' if i%2 else None)

doc.add_paragraph()

# =========================================================
#  2. DANH SACH LOP VA QUAN HE
# =========================================================
heading(doc, '2. DANH SACH LOP VA QUAN HE', 1)
heading(doc, '2.1 Bang tong hop cac lop (theo code thuc te)', 2)
para(doc, 'Cac lop duoi day khop chinh xac voi cac file .java trong package com.duong.salesmanagement.model')

cl = doc.add_table(rows=1, cols=4)
cl.style = 'Table Grid'
tbl_header(cl, ['Ten lop (file .java)', 'Stereotype', 'Bang DB', 'Thuoc tinh -- Phuong thuc chinh'])

# Dữ liệu khớp 100% với code thực tế đã đọc
class_data = [
    ('User',
     '<<entity>>',
     'users',
     '- id: Long\n- username: String (unique, not null)\n- password: String (not null)\n- fullName: String (not null)\n- role: Role (enum, not null)\n---\n(getter / setter)'),
    ('CustomerProfile',
     '<<entity>>',
     'customer_profiles',
     '- id: Long\n- user: User  [@OneToOne, CascadeALL]\n- phoneNumber: String\n- deliveryAddress: String\n---\n(getter / setter)'),
    ('RestaurantProfile',
     '<<entity>>',
     'restaurant_profiles',
     '- id: Long\n- user: User  [@OneToOne, CascadeALL]\n- restaurantName: String\n- address: String\n- isOpen: boolean\n- averageRating: Double\n---\n+ isOpen(): boolean'),
    ('DriverProfile',
     '<<entity>>',
     'driver_profiles',
     '- id: Long\n- user: User  [@OneToOne, CascadeALL]\n- licensePlate: String\n- phoneNumber: String\n- isAvailable: boolean\n---\n(getter / setter)'),
    ('Category',
     '<<entity>>',
     'categories',
     '- id: Long\n- name: String (unique, not null)\n- description: String\n---\n(getter / setter)'),
    ('MenuItem',
     '<<entity>>',
     'menu_items',
     '- id: Long\n- restaurant: RestaurantProfile  [@ManyToOne, LAZY]\n- category: Category  [@ManyToOne, LAZY]\n- name: String\n- description: String\n- price: Double\n- imageUrl: String\n- isAvailable: boolean\n---\n+ isAvailable(): boolean'),
    ('FoodOrder',
     '<<entity>>',
     'food_orders',
     '- id: Long\n- customer: CustomerProfile  [@ManyToOne, LAZY]\n- restaurant: RestaurantProfile  [@ManyToOne, LAZY]\n- driver: DriverProfile  [@ManyToOne, LAZY]\n- status: OrderStatus  [@Enumerated]\n- totalAmount: Double\n- orderTime: LocalDateTime\n- deliveryAddress: String\n- orderItems: List<OrderItem>  [@OneToMany, CascadeALL]\n---\n(getter / setter)'),
    ('OrderItem',
     '<<entity>>',
     'order_items',
     '- id: Long\n- order: FoodOrder  [@ManyToOne, LAZY]\n- menuItem: MenuItem  [@ManyToOne, LAZY]\n- quantity: Integer\n- priceAtTimeOfOrder: Double\n---\n+ getSubtotal(): Double'),
    ('Payment',
     '<<entity>>',
     'payments',
     '- id: Long\n- order: FoodOrder  [@OneToOne, LAZY]\n- paymentMethod: PaymentMethod  [@Enumerated]\n- paymentStatus: PaymentStatus  [@Enumerated]\n- amount: Double\n- transactionDate: LocalDateTime\n---\n(getter / setter)'),
    ('Review',
     '<<entity>>',
     'reviews',
     '- id: Long\n- order: FoodOrder  [@OneToOne, LAZY]\n- rating: Integer\n- comment: String\n- createdAt: LocalDateTime\n---\n(getter / setter)'),
    ('Voucher',
     '<<entity>>',
     'vouchers',
     '- id: Long\n- code: String (unique, not null)\n- discountValue: Double\n- discountType: DiscountType  [@Enumerated]\n- expirationDate: LocalDate\n- isActive: boolean\n---\n+ isActive(): boolean'),
]
for i, rd in enumerate(class_data):
    fill_row(cl.add_row(), rd, alt='EBF3FB' if i%2 else None)

doc.add_paragraph()
heading(doc, '2.2 Cac lop Enum', 2)
para(doc, 'Cac enum duoc dinh nghia trong package model, su dung @Enumerated(EnumType.STRING):')

et = doc.add_table(rows=1, cols=3)
et.style = 'Table Grid'
tbl_header(et, ['Ten Enum (file .java)', 'Cac gia tri', 'Su dung trong lop nao'])
enum_data = [
    ('Role',          'CUSTOMER, RESTAURANT, DRIVER, ADMIN',       'User.role'),
    ('OrderStatus',   'PENDING, CONFIRMED, PREPARING, DELIVERING,\nDELIVERED, CANCELLED', 'FoodOrder.status'),
    ('PaymentMethod', 'CASH, VNPAY, MOMO',                         'Payment.paymentMethod'),
    ('PaymentStatus', 'PENDING, COMPLETED, FAILED, REFUNDED',      'Payment.paymentStatus'),
    ('DiscountType',  'PERCENTAGE, FIXED_AMOUNT',                  'Voucher.discountType'),
]
for i, rd in enumerate(enum_data):
    fill_row(et.add_row(), rd, alt='FFF2CC' if i%2 else None)

doc.add_paragraph()
heading(doc, '2.3 Bang quan he giua cac lop', 2)
para(doc, 'Quan he khop chinh xac voi annotation JPA trong code thuc te.')

rt = doc.add_table(rows=1, cols=5)
rt.style = 'Table Grid'
tbl_header(rt, ['Lop A (chu)', 'Lop B (phu)', 'Loai quan he', 'Multiplicity', 'JPA Annotation / Ghi chu'])

# Quan hệ thực tế 100% từ code đã đọc
rel_data = [
    # User <-> Profiles (Composition via OneToOne)
    ('CustomerProfile', 'User',
     'Composition (1 chieu)',
     '1 : 1',
     '@OneToOne(cascade=ALL)\n@JoinColumn(user_id)\nCustomerProfile so huu User'),
    ('RestaurantProfile', 'User',
     'Composition (1 chieu)',
     '1 : 1',
     '@OneToOne(cascade=ALL)\n@JoinColumn(user_id)\nRestaurantProfile so huu User'),
    ('DriverProfile', 'User',
     'Composition (1 chieu)',
     '1 : 1',
     '@OneToOne(cascade=ALL)\n@JoinColumn(user_id)\nDriverProfile so huu User'),
    # MenuItem
    ('RestaurantProfile', 'MenuItem',
     'Association (1 chieu)',
     '1 : 0..*',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(restaurant_id)\nMenuItem biet RestaurantProfile cua no'),
    ('Category', 'MenuItem',
     'Association (1 chieu)',
     '1 : 0..*',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(category_id)\nMenuItem biet Category cua no'),
    # FoodOrder
    ('FoodOrder', 'CustomerProfile',
     'Association (nhieu phia)',
     '0..* : 1',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(customer_id, not null)'),
    ('FoodOrder', 'RestaurantProfile',
     'Association (nhieu phia)',
     '0..* : 1',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(restaurant_id, not null)'),
    ('FoodOrder', 'DriverProfile',
     'Association (nhieu phia)',
     '0..* : 0..1',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(driver_id)\nnullable=true - co the chua co driver'),
    ('FoodOrder', 'OrderStatus',
     'Dependency (uses)',
     '1 : 1',
     '@Enumerated(EnumType.STRING)\nFoodOrder.status'),
    # OrderItem (Composition cua FoodOrder)
    ('FoodOrder', 'OrderItem',
     'Composition',
     '1 : 1..*',
     '@OneToMany(mappedBy="order", cascade=ALL)\nOrderItem bi xoa khi FoodOrder bi xoa'),
    ('OrderItem', 'FoodOrder',
     'Association (nguoc)',
     '0..* : 1',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(order_id, not null)'),
    ('OrderItem', 'MenuItem',
     'Association (snapshot)',
     '0..* : 1',
     '@ManyToOne(fetch=LAZY)\n@JoinColumn(menu_item_id, not null)\nLuu gia tai thoi diem dat hang'),
    # Payment
    ('Payment', 'FoodOrder',
     'Association (1 chieu)',
     '1 : 1',
     '@OneToOne(fetch=LAZY)\n@JoinColumn(order_id, unique, not null)'),
    ('Payment', 'PaymentMethod',
     'Dependency (uses)',
     '1 : 1',
     '@Enumerated(EnumType.STRING)'),
    ('Payment', 'PaymentStatus',
     'Dependency (uses)',
     '1 : 1',
     '@Enumerated(EnumType.STRING)'),
    # Review
    ('Review', 'FoodOrder',
     'Association (1 chieu)',
     '1 : 1',
     '@OneToOne(fetch=LAZY)\n@JoinColumn(order_id, unique, not null)'),
    # Voucher (doc lap, khong co FK trong code hien tai)
    ('Voucher', 'DiscountType',
     'Dependency (uses)',
     '1 : 1',
     '@Enumerated(EnumType.STRING)'),
    # User enum
    ('User', 'Role',
     'Dependency (uses)',
     '1 : 1',
     '@Enumerated(EnumType.STRING)\nUser.role'),
]
for i, rd in enumerate(rel_data):
    fill_row(rt.add_row(), rd, alt='E8EAF6' if i%2 else None)

doc.add_paragraph()
para(doc, 'Tong so: 18 quan he (3 composition, 9 association, 6 dependency voi enum)', italic=True)
doc.add_paragraph()

# =========================================================
#  3. CAU TRUC DU AN
# =========================================================
heading(doc, '3. CAU TRUC DU AN JAVA SPRING BOOT', 1)
heading(doc, '3.1 Mo hinh kien truc phan tang', 2)

at = doc.add_table(rows=1, cols=3)
at.style = 'Table Grid'
tbl_header(at, ['Tang', 'Package', 'Vai tro'])
arch = [
    ('Presentation',  'controller/', 'Nhan HTTP request, goi Service, tra response'),
    ('Business Logic','service/',    'Xu ly nghiep vu: dat hang, tinh tien, kiem tra voucher'),
    ('Data Access',   'repository/', 'Spring Data JPA interfaces - truy van MySQL'),
    ('Domain',        'model/',      'JPA Entities (11 lop) + Enums (5 lop)'),
    ('Cross-cutting', 'config/, security/', 'Spring Security, JWT config'),
]
for i, rd in enumerate(arch):
    fill_row(at.add_row(), rd, alt='E8F5E9' if i%2 else None)

doc.add_paragraph()
heading(doc, '3.2 Cau truc thu muc thuc te', 2)

tree = (
    "gs-serving-web-content-main/complete/\n"
    "pom.xml\n"
    "src/\n"
    "  main/\n"
    "    java/com/duong/salesmanagement/\n"
    "      model/\n"
    "        User.java\n"
    "        CustomerProfile.java\n"
    "        RestaurantProfile.java\n"
    "        DriverProfile.java\n"
    "        Category.java\n"
    "        MenuItem.java\n"
    "        FoodOrder.java\n"
    "        OrderItem.java\n"
    "        Payment.java\n"
    "        Review.java\n"
    "        Voucher.java\n"
    "        Role.java              (enum)\n"
    "        OrderStatus.java       (enum)\n"
    "        PaymentMethod.java     (enum)\n"
    "        PaymentStatus.java     (enum)\n"
    "        DiscountType.java      (enum)\n"
    "      repository/\n"
    "      service/\n"
    "      controller/\n"
    "      config/\n"
    "    resources/\n"
    "      application.properties\n"
    "  test/"
)
p = doc.add_paragraph()
r = p.add_run(tree)
r.font.name = 'Courier New'
r.font.size = Pt(9)
p.paragraph_format.left_indent = Inches(0.3)
doc.add_paragraph()

heading(doc, '3.3 Bang file code skeleton', 2)
ft = doc.add_table(rows=1, cols=2)
ft.style = 'Table Grid'
tbl_header(ft, ['File .java', 'Mo ta'])
files = [
    ('pom.xml',                           'Spring Boot 3.3.0, Java 17, MySQL Connector, Spring Security, JWT'),
    ('application.properties',            'DB=Aiven MySQL, port=8080, JPA=update, Thymeleaf cache=false'),
    ('model/User.java',                   '@Entity, bang users: id, username, password, fullName, role(Role)'),
    ('model/CustomerProfile.java',        '@Entity, bang customer_profiles: id, user(@OneToOne), phoneNumber, deliveryAddress'),
    ('model/RestaurantProfile.java',      '@Entity, bang restaurant_profiles: id, user(@OneToOne), restaurantName, address, isOpen, averageRating'),
    ('model/DriverProfile.java',          '@Entity, bang driver_profiles: id, user(@OneToOne), licensePlate, phoneNumber, isAvailable'),
    ('model/Category.java',               '@Entity, bang categories: id, name(unique), description'),
    ('model/MenuItem.java',               '@Entity, bang menu_items: id, restaurant(@ManyToOne), category(@ManyToOne), name, price, isAvailable'),
    ('model/FoodOrder.java',              '@Entity, bang food_orders: id, customer/restaurant/driver(@ManyToOne), status(OrderStatus), totalAmount, orderTime, orderItems(@OneToMany)'),
    ('model/OrderItem.java',              '@Entity, bang order_items: id, order(@ManyToOne), menuItem(@ManyToOne), quantity, priceAtTimeOfOrder'),
    ('model/Payment.java',                '@Entity, bang payments: id, order(@OneToOne), paymentMethod, paymentStatus, amount, transactionDate'),
    ('model/Review.java',                 '@Entity, bang reviews: id, order(@OneToOne), rating, comment, createdAt'),
    ('model/Voucher.java',                '@Entity, bang vouchers: id, code(unique), discountValue, discountType(DiscountType), expirationDate, isActive'),
    ('model/Role.java',                   'enum: CUSTOMER, RESTAURANT, DRIVER, ADMIN'),
    ('model/OrderStatus.java',            'enum: PENDING, CONFIRMED, PREPARING, DELIVERING, DELIVERED, CANCELLED'),
    ('model/PaymentMethod.java',          'enum: CASH, VNPAY, MOMO'),
    ('model/PaymentStatus.java',          'enum: PENDING, COMPLETED, FAILED, REFUNDED'),
    ('model/DiscountType.java',           'enum: PERCENTAGE, FIXED_AMOUNT'),
]
for i, rd in enumerate(files):
    fill_row(ft.add_row(), rd, alt='F9F9F9' if i%2 else None)

doc.add_paragraph()

# =========================================================
#  4. GITHUB
# =========================================================
heading(doc, '4. HUONG DAN DAY CODE LEN GITHUB', 1)
heading(doc, '4.1 Lenh Git khoi tao', 2)

gt = doc.add_table(rows=1, cols=2)
gt.style = 'Table Grid'
tbl_header(gt, ['#', 'Lenh terminal'])
git_data = [
    ('1', 'cd gs-serving-web-content-main/complete\ngit init\ngit remote add origin https://github.com/<USERNAME>/PTTKPM25_26_NXX_Nhom12.git'),
    ('2', 'git add .\ngit commit -m "feat(Tuan3): Add 11 entity classes, 5 enums, JPA relationships"'),
    ('3', 'git push -u origin main'),
]
for i, rd in enumerate(git_data):
    row = gt.add_row()
    fill_row(row, rd)
    for rn in row.cells[1].paragraphs[0].runs:
        rn.font.name='Courier New'; rn.font.size=Pt(9)

doc.add_paragraph()
heading(doc, '4.2 Quy uoc commit', 2)
commits = [
    'feat(Tuan3): Add User entity with Role enum',
    'feat(Tuan3): Add CustomerProfile, RestaurantProfile, DriverProfile with OneToOne mapping',
    'feat(Tuan3): Add Category and MenuItem entities',
    'feat(Tuan3): Add FoodOrder and OrderItem with OneToMany mapping',
    'feat(Tuan3): Add Payment and Review with OneToOne mapping',
    'feat(Tuan3): Add Voucher entity with DiscountType enum',
    'docs(Tuan3): Add class diagram and README',
]
for c in commits:
    p = doc.add_paragraph(style='List Paragraph')
    rn = p.add_run(c)
    rn.font.name='Courier New'; rn.font.size=Pt(10)

doc.add_paragraph()

# =========================================================
#  5. CHECKLIST
# =========================================================
heading(doc, '5. CHECKLIST SAN PHAM TUAN 3', 1)
cht = doc.add_table(rows=1, cols=4)
cht.style = 'Table Grid'
tbl_header(cht, ['#', 'San pham', 'Trang thai', 'File / Vi tri'])
checklist = [
    ('1', 'Bang Noun Extraction (11 Entity + 5 Enum)',      'Hoan thanh', 'Muc 1'),
    ('2', 'Bang tong hop 11 lop + thuoc tinh day du',       'Hoan thanh', 'Muc 2.1'),
    ('3', 'Bang 5 lop Enum',                                'Hoan thanh', 'Muc 2.2'),
    ('4', 'Bang 18 quan he voi JPA annotation',             'Hoan thanh', 'Muc 2.3'),
    ('5', 'Cay thu muc du an thuc te',                      'Hoan thanh', 'Muc 3.2'),
    ('6', 'Bang 18 file code skeleton voi mo ta',           'Hoan thanh', 'Muc 3.3'),
    ('7', 'Class Diagram Draw.io (11 lop + 5 enum)',        'Can ve lai', 'ClassDiagram_Team12.drawio'),
    ('8', 'GitHub push code',                               'Can thuc hien', 'Xem Muc 4'),
    ('9', 'Bien ban hop nhom Tuan 3',                       'Can bo sung', 'BienBanHop_Tuan3_Team12.docx'),
]
for i, rd in enumerate(checklist):
    fill_row(cht.add_row(), rd, alt='FEF9E7' if i%2 else None)

out = r'd:\PTTKPM_12\files\ClassDiagram_Tuan3_Team12_Fixed.docx'
doc.save(out)
print('Done! Saved: ' + out)
